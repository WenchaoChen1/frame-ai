"""
DOCX 文档加载器
"""
from typing import BinaryIO

from .base import DocumentLoader
from app.core.logger import get_logger

logger = get_logger(__name__)


class DocxLoader(DocumentLoader):
    """DOCX 文档加载器"""
    
    def extract_text(self, file: BinaryIO) -> str:
        """
        从 DOCX 文件中提取文本
        
        Args:
            file: 文件对象
            
        Returns:
            提取的文本内容
        """
        try:
            from docx import Document
            
            file.seek(0)
            doc = Document(file)
            
            # 提取所有段落的文本
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # 提取表格中的文本
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        text_parts.append(row_text)
            
            logger.info(f"成功提取 DOCX 文本，共 {len(text_parts)} 个段落/行")
            return "\n\n".join(text_parts).strip()
            
        except ImportError:
            logger.error("python-docx 未安装，无法处理 DOCX 文件")
            raise ImportError("需要安装 python-docx 来处理 DOCX 文件")
        except Exception as e:
            logger.error(f"提取 DOCX 文本失败: {e}")
            raise
    
    def supports_file_type(self, file_type: str) -> bool:
        """检查是否支持该文件类型"""
        return file_type.lower() in ['.docx', 'docx']

